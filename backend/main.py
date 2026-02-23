import os
import shutil
import tempfile
import time
import uuid
from pathlib import Path
from typing import Optional, Dict

import docx
import pymupdf4llm
import arxiv
import requests
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse

from openai import OpenAI

# --- ENV & OPENAI SETUP ---
env_path = Path(__file__).parent / ".env"
load_dotenv(dotenv_path=env_path)

API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise RuntimeError("CRITICAL ERROR: OPENAI_API_KEY not found in .env file.")

client = OpenAI(api_key=API_KEY)

FALLBACK_MODEL = "gpt-4o-mini"

app = FastAPI(title="Next-Gen APA Generator with BibTeX")

# In-memory session store: session_id -> {"docx_path": ..., "bib_path": ...}
SESSION_STORE: Dict[str, Dict[str, str]] = {}


def extract_markdown_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF as Markdown.
    For performance, we limit to the first 2 pages by default.
    """
    try:
        return pymupdf4llm.to_markdown(file_path, pages=[0, 1])
    except Exception as e:
        print(f"PDF read error: {e}")
        return ""


def generate_citation_with_retry(combined_text: str, model_id: str, retries: int = 2) -> str:
    """
    Generate APA 7 citations from combined text using the requested model.
    On certain errors, fall back to a lighter model.
    """
    prompt = (
        "You are an expert academic writing assistant.\n"
        "You will receive excerpts from one or more academic papers.\n"
        "Your task is to create a bibliography entry for EACH distinct paper in APA 7 format.\n\n"
        "Requirements:\n"
        "- Output ONLY the list of citations, one citation per line.\n"
        "- Do NOT include any introductory or concluding text.\n"
        "- If some bibliographic fields are missing, do NOT invent them.\n"
        "- Use 'n.d.' for missing years and 'Unknown' where appropriate.\n"
        "- If you can infer a DOI or arXiv ID from the text, include it; otherwise omit it.\n"
        "- Sort the citations alphabetically by first author's surname if possible.\n\n"
        "Source excerpts:\n"
        f"{combined_text}"
    )

    print(f"🤖 Sending request to model: {model_id}")
    current_model = model_id

    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model=current_model,
                messages=[
                    {"role": "system", "content": "You format academic references in APA 7 style."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
            )
            text = response.choices[0].message.content or ""
            return text.strip()

        except Exception as e:
            error_msg = str(e)
            print(f"⚠️ Error ({current_model}): {error_msg}")

            if "404" in error_msg or "not found" in error_msg.lower() or "403" in error_msg:
                if current_model != FALLBACK_MODEL:
                    print(f"🔄 Model not available. Falling back to '{FALLBACK_MODEL}'...")
                    current_model = FALLBACK_MODEL
                    continue
            elif "429" in error_msg or "rate limit" in error_msg.lower():
                wait_time = 10
                print(f"⏳ Rate limit hit. Waiting {wait_time} seconds...")
                time.sleep(wait_time)
                if attempt == retries - 1 and current_model != FALLBACK_MODEL:
                    print(f"🔄 Still limited, trying lighter model '{FALLBACK_MODEL}'.")
                    current_model = FALLBACK_MODEL
                    continue
            else:
                raise HTTPException(status_code=500, detail=f"LLM Error: {error_msg}")

    raise HTTPException(status_code=429, detail="Service is temporarily unavailable, please try again later.")


def generate_bibtex_from_apa(apa_citations: str, model_id: str) -> str:
    """
    Convert APA 7 citations into BibTeX using the LLM.
    """
    prompt = (
        "Convert the following APA 7 citations into valid BibTeX entries.\n"
        "Requirements:\n"
        "- Output ONLY BibTeX entries, no explanations.\n"
        "- Use stable citation keys (e.g., firstauthor_year_titleword).\n"
        "- Preserve DOIs, URLs, and arXiv IDs when present.\n\n"
        f"{apa_citations}"
    )

    response = client.chat.completions.create(
        model=model_id,
        messages=[
            {"role": "system", "content": "You convert APA citations to valid BibTeX entries."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
    )
    text = response.choices[0].message.content or ""
    return text.strip()


def add_arxiv_papers_to_corpus(
    temp_dir: str,
    combined_prompt_text: str,
    arxiv_query: Optional[str],
    arxiv_max_results: int,
) -> str:
    """
    Search arXiv for papers and add their PDFs to the corpus.
    The PDFs are downloaded into temp_dir and processed like uploaded files.
    """
    if not arxiv_query or arxiv_max_results <= 0:
        return combined_prompt_text

    try:
        search = arxiv.Search(
            query=arxiv_query,
            max_results=arxiv_max_results,
            sort_by=arxiv.SortCriterion.Relevance,
        )

        for result in search.results():
            pdf_path = os.path.join(temp_dir, f"{result.get_short_id()}.pdf")
            try:
                pdf_url = result.pdf_url
                r = requests.get(pdf_url, timeout=60)
                r.raise_for_status()
                with open(pdf_path, "wb") as f:
                    f.write(r.content)

                text = extract_markdown_from_pdf(pdf_path)
                if text:
                    meta_header = (
                        f"--- SOURCE: arXiv {result.get_short_id()} ---\n"
                        f"Title: {result.title}\n"
                        f"Authors: {', '.join(a.name for a in result.authors)}\n"
                        f"arXiv ID: {result.get_short_id()}\n"
                        f"URL: {result.entry_id}\n\n"
                    )
                    combined_prompt_text += meta_header + text[:5000] + "\n\n"
            except Exception as e:
                print(f"Error downloading or processing arXiv paper {result.get_short_id()}: {e}")

    except Exception as e:
        print(f"arXiv search error: {e}")

    return combined_prompt_text


@app.post("/prepare-bibliography/")
async def prepare_bibliography(
    files: list[UploadFile] = File(...),
    model_id: str = Form(...),
    arxiv_query: Optional[str] = Form(None),
    arxiv_max_results: int = Form(0),
):
    """
    Prepare bibliography once:
    - Process uploaded PDFs
    - Optionally fetch and add arXiv papers
    - Generate APA 7 citations
    - Generate BibTeX from APA
    - Store both as temp files
    - Return a session_id for later download
    """
    temp_dir = tempfile.mkdtemp()

    try:
        combined_prompt_text = ""

        # 1. Process uploaded PDF files
        for file in files:
            temp_file_path = os.path.join(temp_dir, file.filename)
            with open(temp_file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            text = extract_markdown_from_pdf(temp_file_path)
            if text:
                combined_prompt_text += f"--- SOURCE: {file.filename} ---\n{text[:5000]}\n\n"

        # 2. Optionally add arXiv papers
        combined_prompt_text = add_arxiv_papers_to_corpus(
            temp_dir=temp_dir,
            combined_prompt_text=combined_prompt_text,
            arxiv_query=arxiv_query,
            arxiv_max_results=arxiv_max_results,
        )

        if not combined_prompt_text:
            shutil.rmtree(temp_dir)
            raise HTTPException(status_code=400, detail="No readable text found in PDFs or arXiv results.")

        # 3. Generate APA citations
        apa_citations = generate_citation_with_retry(combined_prompt_text, model_id)

        # 4. Generate BibTeX from APA
        bibtex_entries = generate_bibtex_from_apa(apa_citations, model_id)

        # 5. Create Word document with APA references
        doc = docx.Document()
        doc.add_heading("References (APA 7)", 0)
        for line in apa_citations.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.add_paragraph(f"\n(Generated with model: {model_id})")

        fd_docx, docx_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd_docx)
        doc.save(docx_path)

        # 6. Create BibTeX file
        fd_bib, bib_path = tempfile.mkstemp(suffix=".bib")
        os.close(fd_bib)
        with open(bib_path, "w", encoding="utf-8") as f:
            f.write(bibtex_entries)

        # 7. Store session
        session_id = str(uuid.uuid4())
        SESSION_STORE[session_id] = {
            "docx_path": docx_path,
            "bib_path": bib_path,
            "temp_dir": temp_dir,
        }

        return JSONResponse({"session_id": session_id})

    except Exception as e:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise e


@app.get("/download/apa/{session_id}")
async def download_apa(session_id: str):
    """
    Download APA 7 Word document for a given session.
    """
    session = SESSION_STORE.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found.")

    docx_path = session.get("docx_path")
    if not docx_path or not os.path.exists(docx_path):
        raise HTTPException(status_code=404, detail="APA file not found.")

    return FileResponse(
        path=docx_path,
        filename="references_apa7.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/download/bibtex/{session_id}")
async def download_bibtex(session_id: str):
    """
    Download BibTeX file for a given session.
    """
    session = SESSION_STORE.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found.")

    bib_path = session.get("bib_path")
    if not bib_path or not os.path.exists(bib_path):
        raise HTTPException(status_code=404, detail="BibTeX file not found.")

    return FileResponse(
        path=bib_path,
        filename="references.bib",
        media_type="application/x-bibtex",
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="127.0.0.1", port=8001)
